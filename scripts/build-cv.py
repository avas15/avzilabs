"""
Generates Avas Saeed's CV as a .docx.

Kept as a script rather than a hand-edited document so the content lives in one
readable place and the formatting cannot drift between revisions. The previous
file had mojibake throughout (en dashes and pound signs written in cp1252 and
read back as UTF-8); building from source avoids that entirely.

    python scripts/build-cv.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# Deliberately NOT under public/: the CV carries a personal phone number and
# this repo is public. Generated into an ignored directory and shared directly.
OUT = Path(__file__).resolve().parents[1] / "private" / "Avas Saeed - CV.docx"

NAME = "AVAS SAEED"
CONTACT = "avas.saeed@gmail.com  |  07521 900531  |  London, UK  |  avzilabs.com  |  github.com/avas15"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
ACCENT = RGBColor(0x0A, 0x30, 0x58)  # prussian, matching the site

SUMMARY = (
    "Strategy and technology lead in telecom spectrum, operating with full autonomy across "
    "auction strategy, valuation and the software that supports both. I have run bid strategy "
    "for five European awards, built the simulation, tracking and live-support tooling those "
    "auctions ran on, and replaced recurring third-party data costs with automated pipelines. "
    "Equally comfortable in the quantitative work (Monte Carlo valuation, DCF/WACC, "
    "combinatorial auction theory) and in shipping the systems (Python, TypeScript, LLM "
    "retrieval, agentic tooling). Imperial College Management Development Programme and London "
    "Business School AI programme."
)

COMPETENCIES = [
    ("Spectrum strategy and auctions",
     "Auction simulation and bid tracking, live auction support, allocation-round reverse "
     "engineering, reserve price modelling, combinatorial winner determination and core "
     "pricing, beauty contest scoring, regulator benchmarking"),
    ("Valuation and financial modelling",
     "Monte Carlo valuation, DCF and WACC, NPV and EBITDA multiples, spectrum price "
     "benchmarking with population and term normalisation, geospatial coverage-based "
     "valuation, sensitivity and scenario analysis"),
    ("AI and machine learning",
     "RAG architectures over regulatory and technical corpora, embedding search and chunking "
     "strategy, local LLM deployment (Qwen and Gemma on llama.cpp) for private data, model "
     "fine-tuning, OCR pipelines, text-to-SQL, text to speech"),
    ("Agentic engineering",
     "Claude Code to an advanced level including custom MCP servers, skills and hooks; MCP "
     "servers built for Unity, Photoshop and FL Studio; OpenCode; harness and context design"),
    ("Data engineering",
     "Python scraping and automation, fully automated daily ETL from regulator and industry "
     "sources, SQL and database design, AWS, GeoPandas and PostGIS, VBA and Excel tooling"),
    ("Software",
     "Python, TypeScript, C#, SQL, PHP, R, JavaScript, HTML/CSS, Git, Cloudflare Workers and "
     "Durable Objects, Unity, Blender, CAD (EzCAD), Bash and PowerShell"),
    ("Leadership and communication",
     "Leading auction strategy workshops and internal training, cross-functional work with "
     "Group Finance and the GSMA, C-suite reporting, full ownership of strategic initiatives"),
]

EXPERIENCE = [
    {
        "role": "Strategy Manager",
        "org": "Orange  |  May 2022 - Present",
        "bullets": [
            "Led end-to-end auction strategy for Orange across Poland (x2), Spain (26 GHz), "
            "Moldova and Slovakia: simulations, bid tracking, pre-auction planning and live "
            "auction support, securing materially better pricing than competitors and saving "
            "tens of millions across the awards.",
            "Reverse-engineered allocation-round mechanics to derive optimal bid and ceiling "
            "strategies, directly shaping outcomes in live auctions.",
            "Built a proprietary Python scraping platform and automated daily ETL across "
            "Cullen International and EU/UK regulator sources, eliminating around GBP 100,000 "
            "per year in third-party data costs and making the internal spectrum database the "
            "team's single source of truth.",
            "Deployed an auction simulation platform and bid-tracking toolkit that reduced "
            "external consultancy spend by over GBP 200,000 in FY2023 and let internal teams "
            "run auctions without outside support.",
            "Built an LLM retrieval system over Orange's spectrum documents and database, with "
            "structure-aware chunking and cited answers, enabling natural-language querying of "
            "holdings, auction rules and pricing benchmarks.",
            "Fine-tuned a model to navigate regulator documentation, accelerating compliance "
            "and due-diligence workflows.",
            "Developed a combinatorial auction lab implementing winner determination and "
            "core-selecting payment rules from the literature (Sandholm; Day and Raghavan; "
            "Ausubel and Milgrom), validated by recreating published experimental conditions "
            "rather than trusting plausible output.",
            "Produced EC D2D spectrum research, beauty contest scoring models and full "
            "valuation models using Monte Carlo methods, WACC, EBITDA multiples and core "
            "pricing algorithms.",
            "Built geospatial valuation tooling for mobile satellite spectrum, replacing "
            "national per-head benchmarks with gridded population, terrain and maritime "
            "coverage analysis where the standard method is systematically wrong.",
            "Delivered ad-hoc benchmarking across MENA markets, establishing pricing and "
            "regulatory comparisons that informed market-entry decisions.",
            "Partnered with Group Finance and the GSMA on holdings data and NPV cost "
            "projections for regulatory and investor reporting.",
            "Led internal training on auction mechanics, valuation methodology and bidding "
            "strategy, upskilling cross-functional teams.",
            "Architected a greenfield auction platform (TypeScript monorepo, Fastify, React, "
            "Postgres): a configuration-driven engine covering SMRA, clock and combinatorial "
            "formats, an auctioneer control room with deterministic round rollback, and a "
            "results pipeline with per-audience visibility.",
        ],
    },
    {
        "role": "Mathematics Teacher",
        "org": "Burnt Mill Academy  |  Sep 2021 - May 2022",
        "bullets": [
            "Led GCSE mathematics classes, achieving a 9% average increase in summative "
            "assessment results.",
            "Designed a coordinated scheme of work with the department and built data-driven "
            "lesson planning from national, local and school-level analysis.",
        ],
    },
    {
        "role": "Summer Intern",
        "org": "Baker McKenzie LLP  |  Aug 2017 - Sep 2017",
        "bullets": [
            "Conducted M&A research analysis and presented findings to senior partners.",
            "Led a cross-functional workshop to design and pitch an application concept at "
            "Google offices.",
        ],
    },
]

EDUCATION = [
    ("BSc Economics with Accounting, First Class Honours", "City, University of London  |  2021"),
    ("Management Development Programme", "Imperial College London  |  2025"),
    ("The Business of AI", "London Business School  |  2025"),
]

PROJECTS = [
    ("Auction Platform",
     "Configuration-driven auction engine with auctioneer console, live bidder screens and "
     "deterministic replay. TypeScript monorepo with the engine and pricing logic as pure, "
     "heavily tested packages."),
    ("SpectrumWDP",
     "Combinatorial winner determination with VCG, bidder-optimal core and nearest-VCG payment "
     "rules, validated against published results from the auction theory literature."),
    ("Spectrum intelligence pipeline",
     "Automated daily ingestion of regulator publications and industry sources with "
     "near-duplicate detection, change detection and structured classification."),
    ("MSS geospatial valuation",
     "Coverage-based valuation for mobile satellite spectrum using gridded population, terrain "
     "and maritime demand instead of national averages."),
    ("Local AI stack",
     "Qwen and Gemma on llama.cpp with RAG, embeddings, OCR and text to speech, so commercially "
     "sensitive material never leaves local hardware."),
    ("MCP toolchain",
     "Model Context Protocol servers exposing Unity, Photoshop and FL Studio as agent-callable "
     "tools, used daily with Claude Code."),
    ("Cat Dog Fish",
     "Real-time multiplayer party game on Cloudflare Workers and Durable Objects, with "
     "authoritative server scoring and a community dictionary decided by a two-thirds vote. "
     "Live at avzilabs.com."),
    ("Imran Football Coaching",
     "Website and custom CRM for a coaching business, integrating scheduling, attendance and "
     "payments with the WhatsApp and Gmail workflows the business already runs on."),
]

INTERESTS = (
    "Padel  -  Game development (Unity, Blender)  -  Cooking  -  Staying on the cusp of AI and "
    "agentic engineering  -  Building local AI tooling"
)


# ----------------------------------------------------------------- formatting

def style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.06

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT
    # Rule under the heading, which reads better than a border-less gap.
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0A3058")
    borders.append(bottom)
    pPr.append(borders)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.13)
    p.add_run(text).font.size = Pt(9.5)


def labelled(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(9.5)
    b = p.add_run(body)
    b.font.size = Pt(9.5)


def build() -> None:
    doc = Document()
    style_base(doc)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(CONTACT)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    heading(doc, "Profile")
    p = doc.add_paragraph()
    p.add_run(SUMMARY).font.size = Pt(9.5)

    heading(doc, "Core competencies")
    for label, body in COMPETENCIES:
        labelled(doc, label, body)

    heading(doc, "Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(job["role"])
        r.bold = True
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(job["org"])
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        r.italic = True

        for b in job["bullets"]:
            bullet(doc, b)

    heading(doc, "Education and professional development")
    for title, org in EDUCATION:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r2 = p.add_run(f"    {org}")
        r2.font.size = Pt(9)
        r2.font.color.rgb = MUTED

    heading(doc, "Selected projects")
    for title, body in PROJECTS:
        labelled(doc, title, body)

    heading(doc, "Interests")
    p = doc.add_paragraph()
    p.add_run(INTERESTS).font.size = Pt(9.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
